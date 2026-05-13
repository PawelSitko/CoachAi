"""
Build user_survey.docx — a participant-ready survey instrument that you can
copy/paste into Google Forms (5 minutes of clicking) and send to testers.

Run:
    python3 build_survey.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

s = doc.styles["Heading 1"]
s.font.name = "Arial"
s.font.size = Pt(16)
s.font.bold = True
s.font.color.rgb = RGBColor(0, 0, 0)

s = doc.styles["Heading 2"]
s.font.name = "Arial"
s.font.size = Pt(13)
s.font.bold = True
s.font.color.rgb = RGBColor(0, 0, 0)

section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)


# ---------- Title ----------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CoachAI: User Evaluation Survey")
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Final Year Project — Pawel Sitko, Goldsmiths University of London")
r.italic = True


# ---------- Participant Information ----------
doc.add_paragraph("Participant Information", style="Heading 1")
doc.add_paragraph(
    "Thank you for taking part in this evaluation. CoachAI is a final year project at "
    "Goldsmiths University of London. It is a web app that generates a weekly strength "
    "and conditioning (S&C) gym plan tailored to your sport, playing position, training "
    "goal and available equipment. The aim of this short survey is to gather your "
    "feedback on the clarity of the generated plan and how sport-specific it feels "
    "compared with the gym workouts you usually do."
)
doc.add_paragraph(
    "The survey takes approximately 3–5 minutes. Your responses are anonymous: no "
    "personally identifying information is collected, and survey data will be used only "
    "for the dissertation that this project produces. There are no foreseeable risks to "
    "participation. You may stop the survey at any time without giving a reason; if you "
    "do, your partial responses will not be used."
)
doc.add_paragraph(
    "If you have questions about the project, please contact me at "
    "psitk001@gold.ac.uk."
)


# ---------- Consent ----------
doc.add_paragraph("Consent", style="Heading 1")
doc.add_paragraph(
    "By proceeding past this section you confirm that you have read the participant "
    "information above, that you are aged 18 or over, that you understand your "
    "participation is voluntary and anonymous, and that you consent to your responses "
    "being used in the dissertation as described."
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Q1.  I have read the participant information and consent to take part.")
r.bold = True
doc.add_paragraph("    ☐  Yes, I consent")
doc.add_paragraph("    ☐  No (please close the survey)")
doc.add_paragraph("    [Form type: single choice; required; if 'No' → end of form]").italic = True


# ---------- Section 1: About you ----------
doc.add_paragraph()
doc.add_paragraph("Section 1: About you", style="Heading 1")
doc.add_paragraph(
    "These questions help us interpret the responses; they are anonymous."
)

# Q2 sport
p = doc.add_paragraph()
r = p.add_run("Q2.  Which sport did you generate a plan for?")
r.bold = True
doc.add_paragraph("    ☐  Volleyball")
doc.add_paragraph("    ☐  Basketball")
doc.add_paragraph("    [Form type: single choice; required]").italic = True

# Q3 experience
p = doc.add_paragraph()
r = p.add_run("Q3.  How would you describe your gym training experience?")
r.bold = True
doc.add_paragraph("    ☐  Beginner (under 6 months consistent training)")
doc.add_paragraph("    ☐  Intermediate (6 months – 2 years)")
doc.add_paragraph("    ☐  Advanced (more than 2 years)")
doc.add_paragraph("    [Form type: single choice; required]").italic = True


# ---------- Section 2: Plan clarity (criterion 4) ----------
doc.add_paragraph()
doc.add_paragraph("Section 2: Plan clarity", style="Heading 1")
doc.add_paragraph(
    "Please look at the plan CoachAI generated for you and answer the following on a "
    "scale of 1 (strongly disagree) to 5 (strongly agree)."
)

# Q4 — clarity (criterion 4 measurement)
p = doc.add_paragraph()
r = p.add_run("Q4.  The plan is clearly structured and easy to follow.")
r.bold = True
doc.add_paragraph("    ◯ 1 (strongly disagree)   ◯ 2   ◯ 3   ◯ 4   ◯ 5 (strongly agree)")
doc.add_paragraph("    [Form type: linear scale 1–5; required.  This question maps to project success criterion 4: target average rating ≥ 4/5.]").italic = True


# ---------- Section 3: Sport-specificity (criterion 5) ----------
doc.add_paragraph()
doc.add_paragraph("Section 3: Sport-specificity", style="Heading 1")

# Q5 — relevance
p = doc.add_paragraph()
r = p.add_run("Q5.  The exercises in the plan feel relevant to my sport.")
r.bold = True
doc.add_paragraph("    ◯ 1 (strongly disagree)   ◯ 2   ◯ 3   ◯ 4   ◯ 5 (strongly agree)")
doc.add_paragraph("    [Form type: linear scale 1–5; required.  This question maps to project success criterion 5.]").italic = True

# Q6 — vs usual workouts
p = doc.add_paragraph()
r = p.add_run("Q6.  This plan feels MORE sport-specific than the gym workouts I usually do.")
r.bold = True
doc.add_paragraph("    ◯ 1 (strongly disagree)   ◯ 2   ◯ 3   ◯ 4   ◯ 5 (strongly agree)")
doc.add_paragraph("    [Form type: linear scale 1–5; required.  This question maps to project success criterion 5: target majority agree or strongly agree.]").italic = True


# ---------- Section 4: Free text ----------
doc.add_paragraph()
doc.add_paragraph("Section 4: Open feedback", style="Heading 1")

p = doc.add_paragraph()
r = p.add_run("Q7.  Was there anything in the plan you would change, remove, or add?")
r.bold = True
doc.add_paragraph("    [Form type: long answer text; not required]").italic = True

p = doc.add_paragraph()
r = p.add_run("Q8.  Would you actually use this plan in the gym? Why or why not?")
r.bold = True
doc.add_paragraph("    [Form type: long answer text; not required]").italic = True

p = doc.add_paragraph()
r = p.add_run("Q9.  Anything else you'd like to share about CoachAI? (optional)")
r.bold = True
doc.add_paragraph("    [Form type: long answer text; not required]").italic = True


# ---------- Closing ----------
doc.add_paragraph()
doc.add_paragraph("Closing message", style="Heading 1")
doc.add_paragraph(
    "Thank you for your time. Your responses will be summarised in the dissertation and "
    "used to plan future improvements to CoachAI."
)


# ---------- Notes for Pawel (won't go into the form) ----------
doc.add_page_break()
doc.add_paragraph("Notes for Pawel (do not include in the form)", style="Heading 1")
doc.add_paragraph(
    "Suggested setup in Google Forms (or similar):"
)
doc.add_paragraph("• Title: CoachAI: User Evaluation Survey", style="List Bullet")
doc.add_paragraph("• Description: short version of the Participant Information block.", style="List Bullet")
doc.add_paragraph("• Section 1 (consent): Q1 as a single-choice question; enable 'go to section based on answer' so 'No' jumps to end.", style="List Bullet")
doc.add_paragraph("• Section 2 (about you): Q2 + Q3 single-choice.", style="List Bullet")
doc.add_paragraph("• Section 3 (Likert): Q4, Q5, Q6 as linear-scale questions, 1–5, with 'strongly disagree' / 'strongly agree' as endpoint labels.", style="List Bullet")
doc.add_paragraph("• Section 4 (free text): Q7, Q8, Q9 as paragraph (long answer) questions.", style="List Bullet")
doc.add_paragraph("• Settings: collect emails OFF, limit to 1 response OFF (so testers can submit on a phone without signing in).", style="List Bullet")
doc.add_paragraph()
doc.add_paragraph(
    "Reporting against success criteria: Q4 → Criterion 4 (target average ≥ 4/5). "
    "Q5 + Q6 → Criterion 5 (target majority agree/strongly agree, i.e. responses of 4 or 5)."
)


# ---------- Save ----------
out = os.path.join(os.path.dirname(__file__), "user_survey.docx")
doc.save(out)
print(f"Wrote {out}")
