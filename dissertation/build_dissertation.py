"""
CoachAI Dissertation builder.

Generates dissertation.docx from the structured content below using python-docx.
Run from this folder:  python3 build_dissertation.py

Re-run after edits to regenerate the .docx. The Python file is the source of truth;
the .docx is a build artefact.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

doc = Document()

# Default font: Arial 11pt, 1.5 line spacing — typical Goldsmiths CS submission style.
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# A4 page, 2.5cm margins
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def set_heading_style(style_name, size_pt, bold=True, space_before=18, space_after=6):
    s = doc.styles[style_name]
    s.font.name = "Arial"
    s.font.size = Pt(size_pt)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)
    s.paragraph_format.keep_with_next = True


set_heading_style("Heading 1", 18, space_before=24, space_after=12)
set_heading_style("Heading 2", 14, space_before=18, space_after=6)
set_heading_style("Heading 3", 12, space_before=12, space_after=6)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_para(text, style=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.page_break_before = True
    return p


def add_h2(text):
    return doc.add_paragraph(text, style="Heading 2")


def add_h3(text):
    return doc.add_paragraph(text, style="Heading 3")


def add_bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def add_number(text):
    return doc.add_paragraph(text, style="List Number")


def add_blank():
    doc.add_paragraph("")


def add_page_break():
    doc.add_page_break()


def add_figure(filename, caption, width_inches=2.5, side_by_side=None):
    """Embed an image from screenshots/ as a figure with caption.
    If side_by_side is provided as a (left_filename, right_filename) tuple, the
    two images are placed in a single-row, two-column table side by side and the
    caption applies to the pair.
    """
    base = os.path.join(os.path.dirname(__file__), "screenshots")
    if side_by_side:
        # Render two images side by side using a 2-cell table
        t = doc.add_table(rows=1, cols=2)
        t.autofit = True
        for i, fn in enumerate(side_by_side):
            cell = t.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            path = os.path.join(base, fn)
            if os.path.exists(path):
                run.add_picture(path, width=Inches(width_inches))
        # Remove table borders for a cleaner look
        for row in t.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for edge in ("top", "left", "bottom", "right"):
                    border = OxmlElement(f"w:{edge}")
                    border.set(qn("w:val"), "nil")
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    else:
        path = os.path.join(base, filename)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if os.path.exists(path):
            run.add_picture(path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)


def add_code(text, language=""):
    """Add a code snippet as a monospace paragraph with subtle indent."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    # add a small spacer paragraph after the code block
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_table_row(table, cells, bold_first_row=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        if bold_first_row:
            run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
    return row


# ---------------------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------------------

add_blank()
add_blank()
add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CoachAI")
r.bold = True
r.font.size = Pt(28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Multi-Sport Strength and Conditioning Recommender System for Amateur Athletes")
r.font.size = Pt(16)
r.italic = True

add_blank()
add_blank()
add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Pawel Sitko").font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Student ID: 33788775").font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("BSc Computing").font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Goldsmiths, University of London").font.size = Pt(11)

add_blank()
add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Final Year Project Dissertation").font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Submitted 18 May 2026").font.size = Pt(11)

add_blank()
add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("[Word count: TBC on completion]").italic = True


# ---------------------------------------------------------------------------
# DECLARATION
# ---------------------------------------------------------------------------

add_h1("Declaration")
add_para(
    "I confirm that the work presented in this dissertation is my own. Where the work of "
    "others has been drawn upon, this has been clearly cited and referenced in accordance "
    "with Goldsmiths’ academic conduct guidelines. No part of this submission has been "
    "submitted for any other degree or qualification."
)
add_blank()
add_para("Pawel Sitko")
add_para("18 May 2026")


# ---------------------------------------------------------------------------
# ABSTRACT
# ---------------------------------------------------------------------------

add_h1("Abstract")
add_para(
    "Amateur and university-level team-sport athletes routinely lack access to qualified "
    "strength and conditioning (S&C) support, and existing artificial intelligence fitness "
    "applications are oriented towards general resistance training or endurance sports rather "
    "than position-specific S&C for indoor team sports such as volleyball and basketball. "
    "This dissertation reports on the design, implementation and evaluation of CoachAI, a "
    "mobile-first web prototype that generates a structured weekly S&C plan from four user "
    "inputs: sport, playing position, training goal and available equipment. The system is "
    "implemented as a Flask backend serving a static HTML, CSS and JavaScript frontend, and "
    "uses a transparent rule-based recommender grounded in published S&C literature "
    "(Ramirez-Campillo et al., 2020; Markovic, 2007; Sheppard and Gabbett, 2009) operating "
    "over a curated database of more than forty exercises tagged by sport, equipment "
    "requirement and movement category."
)
add_para(
    "The prototype was evaluated against seven success criteria defined in the project brief. "
    "Functional testing across a matrix of input combinations confirmed that valid, "
    "well-structured plans are produced for both supported sports (volleyball and basketball) "
    "across all three equipment levels (full gym, dumbbells, bodyweight). A small user study "
    "with amateur athletes assessed perceived clarity and sport-specificity of the generated "
    "plans. The dissertation reflects honestly on the deliberate decision to narrow scope from "
    "the original design (which included authentication, a relational database and a hybrid "
    "language-model layer) towards a leaner, fully transparent rule-based system, and "
    "discusses the trade-offs this introduced. Findings inform a focused future-work agenda "
    "centred on additional sports, progression modelling and longitudinal evaluation."
)


# ---------------------------------------------------------------------------
# TABLE OF CONTENTS (manual, since python-docx TOC requires Word to render)
# ---------------------------------------------------------------------------

add_h1("Table of Contents")
toc_entries = [
    ("1.  Introduction", ""),
    ("       1.1  Motivation", ""),
    ("       1.2  Aim and Objectives", ""),
    ("       1.3  Contribution", ""),
    ("       1.4  Structure of the Report", ""),
    ("2.  Background and Literature Review", ""),
    ("       2.1  Strength and Conditioning in Amateur Sport", ""),
    ("       2.2  Position-Specific Demands in Volleyball and Basketball", ""),
    ("       2.3  The State of AI Fitness Applications", ""),
    ("       2.4  Recommender System Architectures", ""),
    ("       2.5  Identified Gap", ""),
    ("3.  Requirements", ""),
    ("       3.1  Stakeholders and User Personas", ""),
    ("       3.2  Functional Requirements", ""),
    ("       3.3  Non-Functional Requirements", ""),
    ("       3.4  Success Criteria", ""),
    ("4.  Design", ""),
    ("       4.1  Cycle 1 Design: Rule-Based Prototype", ""),
    ("       4.2  Cycle 2 Design: Extensions Driven by Phase 1 Findings", ""),
    ("       4.3  Remaining Exclusions", ""),
    ("5.  Implementation", ""),
    ("       5.1  Project Structure", ""),
    ("       5.2  Backend: Flask Application", ""),
    ("       5.3  Backend: The Rule-Based Recommender", ""),
    ("       5.4  Frontend: Mobile-First Web Interface", ""),
    ("       5.5  Exercise Database and Extensibility", ""),
    ("       5.6  Cycle 2: Authentication and Per-User Plan Persistence", ""),
    ("       5.7  Cycle 2: LLM-Augmented Beginner-Friendly Notes", ""),
    ("6.  Evaluation", ""),
    ("       6.1  Evaluation Method", ""),
    ("       6.2  Functional Evaluation (Criteria 1, 2 and 3)", ""),
    ("       6.3  Phase 1: Rule-Based Prototype Evaluation", ""),
    ("       6.4  Phase 2: Extended System Pilot", ""),
    ("       6.5  Documentation and Future-Work Evidence", ""),
    ("       6.6  Threats to Validity", ""),
    ("       6.7  Summary Against the Success Criteria", ""),
    ("7.  Discussion and Reflection", ""),
    ("       7.1  Where CoachAI Performs Well", ""),
    ("       7.2  Where CoachAI Falls Short", ""),
    ("       7.3  Reflecting on the Iterative Process", ""),
    ("       7.4  Methodological Reflection", ""),
    ("8.  Conclusion and Future Work", ""),
    ("       8.1  Summary of Contribution", ""),
    ("       8.2  Future Work", ""),
    ("       8.3  Closing Remarks", ""),
    ("References", ""),
    ("Appendices", ""),
    ("       Appendix A:  Functional Test Matrix", ""),
    ("       Appendix B:  User Survey Instrument", ""),
    ("       Appendix C:  Selected Code Listings", ""),
]
for label, note in toc_entries:
    p = doc.add_paragraph()
    p.add_run(label)
    if note:
        r = p.add_run("   " + note)
        r.italic = True


# ---------------------------------------------------------------------------
# CHAPTER 1: INTRODUCTION  (~500 words)
# ---------------------------------------------------------------------------

add_h1("1.  Introduction")

add_h2("1.1  Motivation")
add_para(
    "University and amateur sports teams in the United Kingdom routinely face performance "
    "barriers that stem from a lack of accessible, sport-specific strength and conditioning "
    "(S&C) guidance. As captain of the Goldsmiths volleyball team, I have direct, repeated "
    "experience of this problem: my teammates and I have struggled to find structured, "
    "evidence-based gym programmes that target the explosive, position-specific qualities "
    "that volleyball demands, in particular vertical jump and lower-body power. Existing "
    "advice tends to come from social media, generic gym templates, or technical coaches "
    "whose expertise lies elsewhere. The result is training that is well-intentioned but "
    "rarely aligned with positional demands or with the published S&C evidence base."
)
add_para(
    "The wider AI fitness landscape does not yet fill this gap. Commercial systems such as "
    "FitnessAI optimise general resistance-training progression (FitnessAI, 2025a), endurance-"
    "oriented platforms such as AI Endurance focus on running, cycling and triathlon (AI "
    "Endurance, 2024a), and recent additions such as Fitbit’s Gemini-powered health coach "
    "(Google/Fitbit, 2025) and Sir Mo Farah’s URUNN running app (Farah and Clarke, 2025) "
    "extend AI coaching into wellness and endurance, respectively. None of these target "
    "position-specific S&C for indoor team sports, and recent commentary cautions that "
    "general AI fitness systems can miss biomechanical and sport-specific loading "
    "considerations that matter for athlete development and injury risk (Ma et al., 2025)."
)

add_h2("1.2  Aim and Objectives")
add_para(
    "The aim of this project is to design, develop and evaluate a prototype of CoachAI, an "
    "AI-powered recommender system that addresses the S&C gap faced by amateur athletes by "
    "generating personalised, evidence-based weekly gym plans from a small set of user inputs "
    "(sport, playing position, training goal and available equipment), with initial support "
    "for volleyball and basketball."
)
add_para("In support of this aim, the project pursues five objectives:")
add_number("Synthesise key S&C principles relevant to volleyball and at least one further sport.")
add_number("Define and document the functional and non-functional requirements for CoachAI.")
add_number("Implement a working mobile-first web prototype that generates structured weekly S&C plans tailored to user profiles.")
add_number("Demonstrate multi-sport capability by supporting plan generation for at least two sports (volleyball and basketball).")
add_number("Conduct an initial user evaluation with amateur athletes to gather evidence on usability, relevance and perceived usefulness of the generated plans.")

add_h2("1.3  Contribution")
add_para(
    "The principal contribution of this work is a working, fully transparent recommender "
    "that maps a four-input user profile to a structured five-day weekly S&C plan, "
    "drawing on a curated exercise database tagged by sport, movement category and "
    "equipment requirement. A secondary contribution is an evaluation framework "
    "structured around seven explicit success criteria, paired with a mixed-methods user "
    "study that surfaces both strengths and limitations of the system in this domain. A "
    "third, methodological contribution is the iterative two-cycle design demonstrated "
    "across the project’s two delivery milestones: a rule-based prototype delivered to "
    "Assignment 3 on 27 April 2026 and evaluated post-submission with amateur athletes; "
    "and an extended final system delivered for Assignment 4, in which authentication, "
    "per-user plan persistence and an LLM-augmented beginner-friendly coaching layer "
    "were added in direct response to specific Phase 1 evaluation findings. The "
    "dissertation reports both cycles and is explicit about which features were present "
    "at each milestone."
)

add_h2("1.4  Structure of the Report")
add_para(
    "Chapter 2 reviews the literature on S&C in amateur sport, position-specific demands in "
    "volleyball and basketball, the state of AI fitness applications and recommender system "
    "architectures, and consolidates the identified gap that motivates CoachAI. Chapter 3 "
    "specifies the requirements, including user personas, functional and non-functional "
    "requirements, and the success criteria against which the project is evaluated. Chapter "
    "4 presents the system design, including the architecture, the recommender model and the "
    "user-experience flows. Chapter 5 describes the implementation, walking through the "
    "Flask backend, the rule-based engine and the curated exercise database. Chapter 6 "
    "reports the evaluation, structured criterion by criterion. Chapter 7 discusses the "
    "results, the limitations and the reflections drawn from the development process. "
    "Chapter 8 concludes and identifies concrete future work grounded in the evaluation."
)


# ---------------------------------------------------------------------------
# CHAPTER 2: BACKGROUND AND LITERATURE REVIEW (~1500 words)
# ---------------------------------------------------------------------------

add_h1("2.  Background and Literature Review")
add_para(
    "This chapter situates CoachAI within four bodies of work: S&C in amateur sport "
    "(Section 2.1); position-specific demands in volleyball and basketball (Section 2.2); "
    "the current state of AI fitness applications (Section 2.3); and recommender system "
    "architectures relevant to health and fitness contexts (Section 2.4). Section 2.5 "
    "consolidates the gap that motivates the prototype."
)

add_h2("2.1  Strength and Conditioning in Amateur Sport")
add_para(
    "Strength and conditioning is widely recognised as a foundational component of athletic "
    "development, particularly in sports characterised by explosive jumping, rapid changes "
    "of direction and repeated high-intensity effort. Meta-analyses consistently show that "
    "plyometric and resistance-based programmes produce reliable improvements in vertical "
    "jump performance and related lower-body power metrics across both youth and adult "
    "athletes (Ma et al., 2025; Markovic, 2007). Newton and Dugan (2002) further argue that "
    "structured strength diagnosis underpins any rational S&C prescription, because the "
    "athlete’s qualitative profile (e.g. relative strength, rate of force development, "
    "reactive strength) determines which adaptations are likely to be limiting."
)
add_para(
    "Despite the strength of this evidence base, access to qualified S&C support is unevenly "
    "distributed. Professional teams typically employ dedicated S&C coaches, but amateur and "
    "university-level athletes frequently rely on social media content, generic gym templates "
    "or advice from technical coaches whose specialism lies elsewhere (Ma et al., 2025). The "
    "consequence is a recurring pattern in amateur sport: athletes who are highly motivated "
    "but who follow training routines that are at best generic and at worst poorly aligned "
    "with their positional demands or long-term development needs."
)
add_para(
    "Two concrete consequences follow from this gap. First, training quality is degraded: "
    "principles such as progressive overload, balanced movement-pattern coverage and "
    "appropriate rest prescription are easy to articulate in coaching textbooks but difficult "
    "to apply consistently without expert oversight, and athletes who self-prescribe often "
    "underweight the qualities (such as posterior-chain strength or reactive strength) that "
    "their sport most rewards. Second, injury risk is elevated: poorly chosen volume or "
    "intensity, particularly for high-impact movements such as plyometric jumps, is "
    "associated with overuse complaints in sports that already place repetitive load on the "
    "lower limb (Markovic, 2007). These consequences frame the problem CoachAI sets out to "
    "address: the goal is not to replace coaching, but to surface evidence-based, sport-"
    "appropriate prescriptions in a form an amateur athlete can act on without specialist "
    "interpretation."
)

add_h2("2.2  Position-Specific Demands in Volleyball and Basketball")
add_para(
    "Volleyball is among the clearest exemplars of a sport in which position dictates "
    "physical demand. Match-analysis research shows that middle blockers accumulate "
    "substantially more high-intensity jumps than other positions because of their repeated "
    "involvement in attacks and block attempts; outside hitters also accumulate significant "
    "jump loads, while setters and liberos perform comparatively more lateral and defensive "
    "movements with fewer maximal jumps (Ramirez-Campillo et al., 2020; Grădinaru, 2021). "
    "These differences are not cosmetic: they have direct implications for training "
    "prescription, since plyometric jump training reliably increases vertical jump "
    "performance across genders and age groups (Ramirez-Campillo et al., 2020), but the "
    "appropriate dose and movement bias differ by position. Sheppard and Gabbett (2009) "
    "argue that, in volleyball, both maximal strength and the ability to express that "
    "strength rapidly (power) are critical, with the relative emphasis varying by role."
)
add_para(
    "Basketball, while not the primary case study of this project, exhibits a comparable "
    "but distinct pattern of positional demand. Guards rely heavily on first-step "
    "acceleration, deceleration and lateral change-of-direction; wings combine sustained "
    "explosiveness with shoulder durability; forwards and centres require contact strength, "
    "raw lower-body force production and shoulder stability for contested rebounds. These "
    "distinctions are well established in coaching practice and are reflected in the "
    "exercise selections and coaching notes implemented in CoachAI’s recommender."
)
add_para(
    "Anecdotal evidence from my own team aligns with these published findings. Middle "
    "blockers report higher rates of lower-body soreness and slower recovery after match "
    "weekends than do liberos or setters, consistent with the elevated jump load implied by "
    "the literature. This convergence of published evidence and lived experience helped "
    "shape one of CoachAI’s core design commitments: that position-specific coaching "
    "notes accompany every plan."
)

add_h2("2.3  The State of AI Fitness Applications")
add_para(
    "AI-driven fitness systems have become commercially significant. FitnessAI uses large "
    "datasets of logged gym sessions to optimise sets, reps and progression for general "
    "resistance training (FitnessAI, 2025a; FitnessAI, 2025b). Endurance-oriented platforms "
    "such as AI Endurance generate adaptive plans for runners, cyclists and triathletes "
    "using GPS, heart-rate and performance data (AI Endurance, 2024a; AI Endurance, 2024b). "
    "More recent entrants extend AI coaching into wellness and recovery: Fitbit’s Gemini-"
    "powered health coach offers personalised recommendations across fitness, recovery and "
    "wellness domains (Google/Fitbit, 2025; Abramson, 2025), while Sir Mo Farah’s URUNN "
    "running app delivers personalised endurance programmes through a mobile platform "
    "(Farah and Clarke, 2025; WithU, 2025)."
)
add_para(
    "Three limitations recur across this landscape. First, none of these systems is built "
    "around position-specific S&C for indoor team sports. A volleyball middle blocker and "
    "a basketball point guard would, on most of these platforms, receive broadly similar "
    "programmes despite radically different performance requirements. Second, several "
    "platforms depend on continuous biometric data from wearables, which restricts their "
    "applicability to athletes who train primarily in indoor gyms without access to "
    "specialised sensors. Third, the recommendation logic is typically opaque to the user: "
    "the system produces a plan but does not explain why a given exercise is included or "
    "how the prescription relates to the user’s stated goal. Beyond functionality, academic "
    "and industry commentary raises explicit concerns about transparency. Generic AI systems "
    "can produce plans that appear sensible but that overlook biomechanical risk factors or "
    "sport-specific loading patterns (Ma et al., 2025; Markovic, 2007). For amateur athletes "
    "training without expert supervision, this opacity is a meaningful risk: the user has no "
    "principled way to discriminate between a recommendation that reflects established S&C "
    "evidence and one that is an artefact of the system’s training distribution."
)

add_h2("2.4  Recommender System Architectures")
add_para(
    "Recommender systems in fitness commonly use rule-based logic, content-based filtering, "
    "or hybrid approaches that combine structured rules with machine-learning outputs "
    "(Ricci, Rokach and Shapira, 2015). Rule-based systems define explicit mappings between "
    "inputs and recommended outputs, which makes them transparent and explainable. Felfernig, "
    "Friedrich and Jannach (2014) note that knowledge-based and rule-based recommenders are "
    "particularly effective in domains where decisions can be grounded in established "
    "expertise and where safety considerations matter, both of which apply to S&C "
    "prescription."
)
add_para(
    "Machine-learning recommenders, by contrast, depend on substantial datasets of user "
    "behaviour, exercises and outcomes. Such datasets are scarce for amateur strength "
    "training, particularly when stratified by sport and position, which makes purely "
    "data-driven approaches impractical at this stage. Hybrid systems combine domain rules "
    "(for example, mapping middle blockers to elevated lower-body power and repeated jumping "
    "work) with AI-assisted generation, often using large language models to structure or "
    "narrate the plan (Cao et al., 2022). This combination has been shown to produce more "
    "flexible yet controlled recommendations in health and fitness contexts."
)
add_para(
    "Three properties of rule-based recommenders make them particularly attractive for "
    "amateur S&C. They are explainable by construction: each output can be traced to a "
    "small set of input mappings, which supports user trust and academic scrutiny. They are "
    "deterministic: the same inputs produce the same outputs, which matters in a domain "
    "where users may revisit and follow a plan over weeks. And they are inexpensive to run "
    "and demonstrate, since no external API or large-scale dataset is required. The "
    "trade-off is reduced flexibility: a rule-based system cannot easily generalise beyond "
    "the prescriptions encoded by its author. CoachAI’s original design anticipated a hybrid "
    "that would mitigate this trade-off by using a language model to vary plan presentation; "
    "the implemented prototype deliberately retained only the rule layer, for reasons "
    "examined in Chapter 4."
)

add_h2("2.5  Identified Gap")
add_para(
    "Two themes emerge consistently from the reviewed literature. First, team-sport athletes, "
    "and volleyball players in particular, benefit from position-specific S&C programming "
    "that respects the differential demands of their roles (Ramirez-Campillo et al., 2020; "
    "Grădinaru, 2021; Sheppard and Gabbett, 2009). Second, current AI fitness platforms "
    "concentrate on general training or endurance sports and do not engage seriously with "
    "the requirements of indoor, court-based team sports."
)
add_para(
    "This convergence reveals a clear and addressable gap: a mobile-first, multi-sport, "
    "position-aware recommender that produces gym-based S&C programmes tailored to the needs "
    "of amateur athletes. CoachAI is the prototype this dissertation develops to determine "
    "whether a transparent, lightweight rule-based system, grounded in the published S&C "
    "evidence base, can deliver more relevant and sport-specific workouts than the generic "
    "alternatives currently available, taking volleyball and basketball as initial case "
    "studies."
)


# ---------------------------------------------------------------------------
# CHAPTER 3: REQUIREMENTS  (~800 words)
# ---------------------------------------------------------------------------

add_h1("3.  Requirements")
add_para(
    "This chapter defines the requirements that drive the design and implementation of "
    "CoachAI. Section 3.1 introduces the stakeholders and user personas that inform design "
    "decisions; Section 3.2 specifies the functional requirements; Section 3.3 captures the "
    "non-functional requirements; and Section 3.4 reproduces the seven success criteria "
    "against which the prototype is evaluated."
)

add_h2("3.1  Stakeholders and User Personas")
add_para(
    "The primary users of CoachAI are amateur and university-level athletes who train "
    "regularly but do not have access to a dedicated S&C coach. The initial user cohort, "
    "and the principal source of evaluation feedback, is the Goldsmiths volleyball team. "
    "Secondary stakeholders include university sports clubs and societies, who could offer "
    "CoachAI as a resource to their members; amateur coaches, who may use it to supplement "
    "technical on-court coaching with structured S&C guidance; and the academic supervisor "
    "and assessors, who evaluate the project’s technical and academic contribution."
)
add_para(
    "Two personas guide design decisions. The first, Woody, is a 21-year-old intermediate-to-"
    "advanced volleyball player whose goals are to improve on-court performance, complement "
    "team training with structured gym work and reduce in-season injury risk. Woody’s "
    "challenges include limited access to qualified S&C support, generic online workouts and "
    "balancing training with academic commitments. The second, Jamie, is a 20-year-old "
    "beginner-to-intermediate basketball player whose goals are general fitness improvement, "
    "confidence in gym-based training and safe progression. Jamie’s challenges include "
    "uncertainty about exercise selection, conflicting online guidance and inconsistent "
    "training frequency around academic deadlines. Together these personas span a useful "
    "range of experience levels and sport-specific needs, and they motivate two design "
    "commitments: clarity of presentation and explicit positional context."
)

add_h2("3.2  Functional Requirements")
add_para(
    "The functional requirements are expressed in natural language and as user stories, "
    "following the convention adopted in the design document. They are grouped by the "
    "two delivery cycles described in Chapter 4: Cycle 1 (FR1–FR8) was delivered for "
    "Assignment 3 on 27 April 2026; Cycle 2 (FR9–FR13) was scoped from the Phase 1 "
    "evaluation findings (Chapter 6) and delivered for Assignment 4."
)
add_h3("Cycle 1 — Core recommender (FR1–FR8)")
add_bullet("FR1: The system shall accept a user profile comprising sport, playing position, training goal and available equipment.")
add_bullet("FR2: The system shall generate a structured weekly S&C plan from a valid user profile.")
add_bullet("FR3: The system shall support at least two sports, namely volleyball and basketball.")
add_bullet("FR4: The system shall present each session as an ordered list of exercises with sets, repetitions, rest intervals and coaching notes.")
add_bullet("FR5: The system shall filter exercises by available equipment, so that users with bodyweight-only access are not shown gym-equipment exercises.")
add_bullet("FR6: The system shall display a position-specific coaching note alongside each generated plan.")
add_bullet("FR7: The system shall allow users to mark sessions as complete and to track weekly progress within a session.")
add_bullet("FR8: The system shall provide an in-app feedback mechanism that captures clarity and sport-specificity ratings.")

add_h3("Cycle 2 — Accounts, persistence and beginner-friendly explanations (FR9–FR13)")
add_bullet("FR9: The system shall allow a user to create an account using email and password, and to sign in and sign out on the same device or another.")
add_bullet("FR10: The system shall allow an authenticated user to save the currently-generated plan, naming it for later retrieval. Saved plans shall be read-only snapshots of their exercise content.")
add_bullet("FR11: The system shall present authenticated users with a list of their saved plans and allow any plan in the list to be re-opened in the plan view.")
add_bullet("FR12: The system shall persist a user’s session-complete state per saved plan, so that mark-complete ticks survive sign-out and are reflected when the plan is re-opened.")
add_bullet("FR13: The system shall offer an optional beginner-friendly explanations mode on the plan view, which replaces each exercise’s coaching note with a plain-language rephrasing authored at development time by a large language model and cached in the project repository.")
add_para(
    "Anonymous use of the recommender is preserved throughout: FR1–FR8 do not require "
    "authentication, and the application’s offline demoability (NFR5) is preserved by "
    "ensuring that all Cycle 2 features are either local to the user’s session or "
    "served from a static cache, never requiring an external API call at runtime.",
    italic=False,
)

add_h3("User stories")
add_bullet("As a student athlete, I want to receive a training plan tailored to my sport and position, so that my gym sessions support my on-court performance.")
add_bullet("As a beginner athlete, I want workouts that are clearly presented and explained, so that I can train safely and confidently.")
add_bullet("As a busy student, I want to generate a plan in under a minute on my phone, so that I can complete profile setup before walking into the gym.")
add_bullet("As a user with limited equipment, I want exercises that match what I actually have available, so that the plan is usable rather than aspirational.")
add_bullet("As a returning user, I want to save a plan I like and re-open it on my next visit, so that I do not have to recreate my profile every time.")
add_bullet("As a novice gym-goer, I want to switch on plain-language exercise instructions, so that an unfamiliar exercise name does not stop me following the plan.")

add_h2("3.3  Non-Functional Requirements")
add_bullet("NFR1 (Usability): the profile-to-plan flow shall fit comfortably on a mobile screen and require no more than three steps before plan generation.")
add_bullet("NFR2 (Performance): plan generation shall complete in under one second on a modest laptop or smartphone over local Wi-Fi, enabling immediate use in a gym setting.")
add_bullet("NFR3 (Determinism): the same input profile shall produce the same plan on every call, so that users can trust and revisit the output rather than chase variation.")
add_bullet("NFR4 (Transparency): every recommended exercise shall include a short note explaining its purpose or technique cue.")
add_bullet("NFR5 (Demoability): the prototype shall run locally without external API keys, persistent storage or internet access beyond loading the page, supporting reliable demonstrations and evaluation sessions.")
add_bullet("NFR6 (Maintainability): adding a new sport, position or exercise shall require only data-level changes (a JSON exercise tag or a new mapping entry), not changes to the core recommender logic.")

add_h2("3.4  Success Criteria")
add_para(
    "The seven success criteria below were defined in the project brief and are reproduced "
    "here verbatim. The Evaluation chapter (Chapter 6) reports on each criterion in turn."
)

# Success criteria table (4 columns: #, Category, Criterion, How measured + target)
table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, t in enumerate(["#", "Category", "Criterion", "Measurement and target"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(10)

criteria = [
    ("1", "Functionality", "User can enter sport, position, goal and equipment and generate a weekly S&C plan.",
     "Manually test different input combinations; verify a valid plan is always produced. Target: 100% of tested inputs produce a valid plan."),
    ("2", "Functionality", "System supports at least two distinct sports (volleyball and basketball).",
     "Generate plans for both sports during testing. Target: plans successfully generated for two or more sports."),
    ("3", "Functionality", "Generated plans are clearly structured (exercises, sets, reps).",
     "Visual inspection by developer and test users. Target: all plans follow a clear, consistent structure."),
    ("4", "User-centred", "Plans are easy to understand and follow.",
     "User survey using a 1-5 rating for clarity. Target: average rating of at least 4 out of 5."),
    ("5", "User-centred", "Users feel plans are sport-specific and more relevant than generic workouts.",
     "User survey and short feedback comparing CoachAI plans to usual routines. Target: majority of users agree or strongly agree."),
    ("6", "Project & research", "System design and implementation are fully documented.",
     "Completion of dissertation sections covering requirements, design and implementation. Target: all core documentation delivered."),
    ("7", "Project & research", "Prototype and feedback provide a clear basis for future work.",
     "Reflection and future-work section grounded in evaluation results. Target: future work grounded in real findings."),
]
for row in criteria:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# CHAPTERS 4-8: PLACEHOLDERS WITH OUTLINES
# ---------------------------------------------------------------------------

add_h1("4.  Design")
add_para(
    "CoachAI was developed using a two-cycle iterative design methodology, in keeping with "
    "the agile and design-thinking framework declared in the project’s design document. "
    "Cycle 1 produced a transparent rule-based prototype focused on the core "
    "recommendation problem; that prototype was evaluated with amateur athletes (Chapter "
    "6, Phase 1), and the findings drove the scope of Cycle 2, which added authentication "
    "and per-user plan persistence and a small LLM-augmented layer for novice-friendly "
    "coaching cues. This chapter documents the design of both cycles. Section 4.1 covers "
    "Cycle 1: the architecture, the recommender model and the user-experience flows of the "
    "rule-based prototype. Section 4.2 covers Cycle 2: the extension that turns the "
    "prototype into the final system. Section 4.3 records the exclusions that remain out "
    "of scope at the end of the project."
)

add_h2("4.1  Cycle 1 Design: Rule-Based Prototype")
add_h3("4.1.1  Architecture")
add_para(
    "Cycle 1 implements a deliberately lean three-component architecture: a Flask "
    "application that serves the static frontend assets, a single POST endpoint that "
    "exposes the rule-based recommender, and a curated JSON exercise database read once "
    "at start-up. There is no relational database, no authentication and no external API "
    "dependency in this cycle. Two considerations drove this choice. First, the prototype "
    "exists to test whether a transparent rule-based recommender can produce credible "
    "position-specific plans for amateur athletes; everything outside that question is "
    "noise that the evaluation does not need. Second, a stateless design preserves "
    "determinism (NFR3) and offline demoability (NFR5) for free, which simplifies "
    "exhaustive functional testing and removes friction from user-evaluation sessions."
)

add_h3("4.1.2  Recommender Model")
add_para(
    "The recommender accepts a four-element user profile (sport, position, goal, equipment) "
    "and returns a structured weekly plan consisting of five sessions, each containing an "
    "ordered list of exercise blocks. Three design choices give the model its shape."
)
add_para(
    "Exercise tagging. Every entry in the exercise database is tagged with five attributes: "
    "a category (e.g. lower_power, plyometric, upper_push, core, agility); an equipment "
    "level (full_gym, dumbbells or bodyweight); a list of sport tags (subset of "
    "{volleyball, basketball}); the principal target tissues; and a short coaching note. "
    "Categories enable the recommender to compose sessions from compatible movement types; "
    "equipment levels enable hierarchical filtering; and sport tags enable a single database "
    "to serve multiple sports without code duplication."
)
add_para(
    "Equipment hierarchy. Equipment is modelled as a strict ordering, full_gym > dumbbells > "
    "bodyweight, encoded as integer ranks. A user with dumbbells can perform any bodyweight "
    "exercise but no barbell movements; a user with full gym access can perform anything. "
    "This produces graceful degradation: bodyweight users never see exercises they cannot "
    "perform, but the same recommender code path serves all three tiers without "
    "special-cased branches."
)
add_para(
    "Session templates. The week is structured as five complementary sessions: "
    "Lower-Body Power, Upper-Body Strength, Lower-Body Strength and Resilience, Power and "
    "Agility, and Core and Recovery. Each template specifies the categories to draw from, "
    "the desired number of exercises and the volume parameters (sets, repetitions, rest), "
    "with values chosen to follow standard S&C prescriptions: three to five sets of four to "
    "eight repetitions for strength work, three to five sets of three to six repetitions "
    "for power and plyometric work, and two to three sets of eight to fifteen repetitions "
    "for accessory and core work. Position context and goal-driven adjustments are layered "
    "on top of these templates: a position note is attached to every plan, and the "
    "vertical_jump goal triggers an additional plyometric block on the lower-power day. "
    "The Power and Agility day branches by sport, prioritising change-of-direction and "
    "lateral plyometrics for basketball and approach-style jump training for volleyball."
)

add_h3("4.1.3  User Experience Flows")
add_para(
    "The Cycle 1 user journey is intentionally short. A landing screen presents the value "
    "proposition and the supported sports. A three-step profile form collects sport and "
    "position, then goal and equipment, then a confirmation summary. On submission the "
    "frontend calls the API and renders the returned plan as a list of session cards, "
    "each containing the session name, focus statement, and exercise blocks with sets, "
    "reps, rest and coaching notes. Sessions can be marked complete, with a small weekly "
    "progress tracker showing how many of the five sessions have been completed; in "
    "Cycle 1 this tracker is in-memory only. The flow is mobile-first throughout: the "
    "application is intended to be opened on a phone in the gym, and the entire "
    "profile-to-plan path can be completed in under a minute on a smartphone over local "
    "Wi-Fi."
)

add_h2("4.2  Cycle 2 Design: Extensions Driven by Phase 1 Findings")
add_para(
    "The Phase 1 evaluation reported in Chapter 6 surfaced two recurring patterns. First, "
    "users explicitly requested the ability to save plans and to track progress across "
    "sessions: three of ten survey respondents asked for some form of account or "
    "progress-tracking functionality, and one cited the absence of tracking as a reason to "
    "use a competing application. Second, novice users in particular found the bare "
    "exercise names and one-line coaching cues insufficient: two of ten respondents asked "
    "for richer instructions or video demonstrations, and the experience-level breakdown "
    "of Q4 ratings showed beginners scoring clarity at 2.25 against advanced users’ 4.50. "
    "Cycle 2 adds two extensions that respond directly to these findings, while preserving "
    "the rule-based core of the recommender unchanged."
)

add_h3("4.2.1  Authentication and Per-User Plan Persistence")
add_para(
    "Cycle 2 introduces an optional account layer. A small SQLite database stores two "
    "entities: a User table that records hashed credentials and a display name, and a "
    "SavedPlan table that holds an immutable JSON snapshot of any plan the user chooses to "
    "save, together with a per-plan record of which sessions the user has marked complete. "
    "Authentication uses Flask-Login session cookies and werkzeug-hashed passwords; no "
    "tokens or third-party identity providers are involved in this cycle. Anonymous use of "
    "the recommender remains a fully supported path: the plan-generation endpoint is open, "
    "the home page does not require sign-in, and the application’s offline demoability "
    "(NFR5) is preserved."
)
add_para(
    "Three design decisions are worth recording. First, saved plans are read-only "
    "snapshots: once a user saves a plan, the underlying exercise selections, volumes and "
    "coaching notes are fixed in the database, even if the recommender is later updated. "
    "This protects the user from silent changes to a plan they have committed to following "
    "and simplifies the evaluation of any later iteration of the recommender. Second, "
    "completion state is stored separately from plan content, so the snapshot remains "
    "immutable while the user’s progress through the snapshot is mutable. Third, the "
    "presence or absence of an account does not change the recommender output: a "
    "logged-in user generating a plan with the same inputs as an anonymous user receives "
    "the same plan, by design, which preserves the deterministic property of NFR3 across "
    "the authenticated and anonymous paths."
)

add_h3("4.2.2  LLM-Augmented Beginner-Friendly Coaching Notes")
add_para(
    "The clarity gap surfaced for novice users is addressed in Cycle 2 by an LLM-"
    "augmented layer that rephrases each exercise’s coaching note into plainer, "
    "beginner-friendly language. The architectural choice that makes this layer safe and "
    "demoable is a strict separation between the LLM, which is consulted only at "
    "development time, and the runtime system, which serves a static cache. A short Python "
    "script (build_tiered_notes.py) iterates over every exercise in the database, prompts "
    "Google Gemini to rephrase the existing coaching note for a beginner audience, and "
    "writes the resulting strings to coaching_notes_tiered.json. The runtime recommender, "
    "consulted on every request, reads this cache and substitutes the beginner phrasing "
    "when the user has activated the corresponding toggle in the user interface; if the "
    "cache is missing it silently falls back to the original coaching note."
)
add_para(
    "Three properties of this design are worth highlighting. The LLM is constrained to "
    "rephrase an existing, expert-authored coaching note rather than to invent new "
    "exercises, modify volumes or generate new technique cues; this constraint is enforced "
    "by the system prompt and reviewed manually by the developer when the cache is "
    "generated. Determinism (NFR3) is preserved because no LLM call is made at runtime: a "
    "given exercise always returns the same beginner phrasing, drawn from a fixed file. "
    "Offline demoability (NFR5) is preserved because the runtime system has no external "
    "dependency once the cache has been built. The trade-off is that any change to a "
    "coaching note requires a developer to re-run the build script; for the exercise "
    "database in this project, this is an acceptable cost."
)

add_h3("4.2.3  User Experience Additions in Cycle 2")
add_para(
    "Cycle 2 introduces three additions to the user interface. The first is an "
    "authentication strip at the top of every screen, which exposes sign-in and create-"
    "account links to anonymous users and a personal greeting, a My Plans link and a sign-"
    "out action to authenticated users. The second is a single primary action button at "
    "the foot of the plan view, labelled Save to my plans, which opens a small modal "
    "where the user names the plan; if the user is not signed in, tapping the button "
    "routes them through the sign-in flow first. The third is a check-box toggle on the "
    "plan view labelled Show beginner-friendly explanations; ticking it re-fetches the "
    "current plan with the beginner_friendly flag set, causing the recommender to swap "
    "each coaching note for the LLM-authored beginner phrasing. The toggle is off by "
    "default so that the rule-based prototype experience is preserved as the baseline. A "
    "My Plans screen lists the user’s saved plans with a per-plan progress indicator (n / "
    "5 sessions completed) and reopens any selected plan into the plan view, restoring "
    "the previously-marked-complete sessions."
)

add_h2("4.3  Remaining Exclusions")
add_para(
    "Several features remain out of scope at the end of Cycle 2. Native iOS and Android "
    "applications were placed out of scope at the project-brief stage and remain so; the "
    "delivered system is a mobile-first responsive web application. Nutrition advice, "
    "wearable integrations and on-court technical drills were similarly excluded by the "
    "brief, and the project focuses solely on gym-based strength and conditioning. The "
    "LLM is used only to rephrase existing coaching notes; it is not used to generate "
    "exercises, modify training volumes or chat with the user, because such uses would "
    "require a substantially deeper safety evaluation than the project timeline allows. "
    "Each of these exclusions is recorded in the future-work agenda of Chapter 8 where "
    "appropriate."
)

add_h1("5.  Implementation")
add_para(
    "This chapter describes the implementation of CoachAI across both cycles. Sections "
    "5.1 to 5.5 cover the Cycle 1 prototype: the project structure, the Flask API, the "
    "rule-based recommender (the technical core), the mobile-first frontend and the "
    "exercise database. Sections 5.6 and 5.7 cover the Cycle 2 extensions: the "
    "authentication and persistence layer, and the LLM-augmented beginner-friendly notes "
    "subsystem."
)

add_h2("5.1  Project Structure")
add_para(
    "The project is organised into a backend and a frontend folder, with no build "
    "tooling, package manager or transpilation step on the frontend side. The "
    "implementation totals approximately one thousand lines of code excluding the JSON "
    "exercise database. The backend is implemented in Python using Flask 3.0.3 as the "
    "only runtime dependency; the frontend uses HTML5, CSS3 and vanilla JavaScript with "
    "no external libraries. This minimal stack supports the demoability requirement "
    "(NFR5): the prototype can be cloned, the dependency installed in a virtual "
    "environment, and the application started with a single command."
)
add_code(
    "coachai/\n"
    "  backend/\n"
    "    app.py            # Flask application + REST endpoint\n"
    "    recommender.py    # Rule-based plan generator\n"
    "    exercises.json    # Curated exercise database (40+ entries)\n"
    "    requirements.txt  # Single dependency: Flask==3.0.3\n"
    "  frontend/\n"
    "    index.html        # Single-page application\n"
    "    styles.css        # Mobile-first responsive stylesheet\n"
    "    app.js            # Profile form, API call, plan rendering"
)

add_h2("5.2  Backend: Flask Application")
add_para(
    "The Flask application in app.py exposes a single REST endpoint, "
    "POST /api/generate-plan, and serves the static frontend assets from the same "
    "process. Request validation is performed in two stages: a lightweight check for the "
    "presence of all required fields in the JSON body, and a deeper validation inside "
    "the recommender that raises ValueError if the sport or equipment is not supported. "
    "Both validation paths return HTTP 400 with a human-readable error message, which the "
    "frontend can display to the user without further translation."
)
add_code(
    "@app.route(\"/api/generate-plan\", methods=[\"POST\"])\n"
    "def api_generate_plan():\n"
    "    data = request.get_json(silent=True) or {}\n"
    "    required = [\"sport\", \"position\", \"goal\", \"equipment\"]\n"
    "    missing = [k for k in required if not data.get(k)]\n"
    "    if missing:\n"
    "        return jsonify({\"error\": f\"Missing fields: {', '.join(missing)}\"}), 400\n"
    "    try:\n"
    "        plan = generate_plan(\n"
    "            sport=data[\"sport\"], position=data[\"position\"],\n"
    "            goal=data[\"goal\"], equipment=data[\"equipment\"],\n"
    "        )\n"
    "    except ValueError as e:\n"
    "        return jsonify({\"error\": str(e)}), 400\n"
    "    return jsonify(plan)"
)
add_para(
    "The application binds to host 0.0.0.0 on port 5001 in development, which makes the "
    "running server reachable from a smartphone on the same Wi-Fi network. This matches "
    "the intended deployment context for evaluation: the developer’s laptop hosts the "
    "Flask process and amateur athletes connect from their own phones. No environment "
    "variables, secrets or external API keys are required, which simplifies "
    "demonstrations and removes a class of failure modes that would otherwise complicate "
    "user-evaluation sessions."
)

add_h2("5.3  Backend: The Rule-Based Recommender")
add_para(
    "The recommender module, recommender.py, is the technical core of CoachAI. It is "
    "structured around three responsibilities: filtering candidate exercises against a "
    "user profile, picking exercises deterministically from a candidate pool, and "
    "composing those picks into the five session templates that make up a weekly plan."
)
add_h3("Equipment hierarchy and filtering")
add_para(
    "Equipment is encoded as a strict ordering. A small dictionary maps each level to an "
    "integer rank, and the helper _allowed returns True when the user’s equipment rank is "
    "at least the rank required by the candidate exercise. This makes the hierarchy "
    "explicit in code and allows the same filter chain to serve users at any level."
)
add_code(
    "EQUIPMENT_RANK = {\"bodyweight\": 0, \"dumbbells\": 1, \"full_gym\": 2}\n"
    "\n"
    "def _allowed(ex: dict, equipment: str) -> bool:\n"
    "    return EQUIPMENT_RANK[ex[\"equipment\"]] <= EQUIPMENT_RANK[equipment]"
)
add_para(
    "The _filter function combines this equipment check with three further constraints: "
    "the exercise must belong to the requested category, must include the user’s sport in "
    "its sport_tags, and must not appear in an optional exclude set. The exclude set is "
    "what enforces uniqueness within a session and across the lower-body sessions, "
    "preventing the recommender from picking the same exercise twice when several "
    "candidates would qualify."
)
add_h3("Deterministic picking")
add_para(
    "Determinism is preserved by sorting the candidate pool on each call before slicing "
    "the first n elements. The sort key is the exercise id, so two equivalent runs of the "
    "recommender produce the same plan in the same order. This satisfies NFR3 without "
    "needing to seed a random number generator or persist any state between calls."
)
add_code(
    "def _pick(category, sport, equipment, n=1, used=None):\n"
    "    pool = _filter(category, sport, equipment, exclude=used)\n"
    "    if not pool:\n"
    "        return []\n"
    "    pool.sort(key=lambda e: e[\"id\"])  # stable, deterministic\n"
    "    picked = pool[:n]\n"
    "    if used is not None:\n"
    "        for ex in picked:\n"
    "            used.add(ex[\"id\"])\n"
    "    return picked"
)
add_h3("Session templates and weekly composition")
add_para(
    "Each of the five session-of-the-week functions follows the same pattern: pick a "
    "small number of exercises from each of the relevant categories, wrap them in blocks "
    "that specify sets, repetitions and rest, and assemble the blocks into a session "
    "object. For example, the lower-body power day pulls two plyometric primers, one main "
    "lower-body lift (preferring lower_power, falling back to lower_strength when no "
    "lower_power option is available at the user’s equipment level), one unilateral "
    "lower-body exercise, and a posterior-chain accessory if the user’s equipment "
    "supports the Nordic hamstring curl. The Power and Agility day branches by sport: "
    "basketball users receive change-of-direction and lateral plyometric work, while "
    "volleyball users receive an approach-jump biased plyometric block."
)
add_para(
    "Goal-driven adjustments are layered on top of the templates after the five sessions "
    "are assembled. The vertical_jump goal triggers an additional plyometric block on the "
    "lower-body power day, drawn from the candidate pool that has not yet been used in "
    "the session. Position context is added independently: every plan carries a "
    "position-specific coaching note from a small lookup table, ensuring that a middle "
    "blocker, a setter and a basketball point guard receive distinct guidance even if "
    "they happen to share the same goal and equipment."
)
add_h3("Output shape")
add_para(
    "The generate_plan function returns a single dictionary containing the user’s inputs, "
    "a focus_summary string derived from the goal, the position_note, and a list of five "
    "session dictionaries. Each session contains a name, a focus statement and a list of "
    "exercise blocks with name, category, sets, reps, rest and notes fields. This shape "
    "is what the frontend consumes directly; no additional transformation occurs between "
    "the API response and the rendered DOM."
)

add_h2("5.4  Frontend: Mobile-First Web Interface")
add_para(
    "The frontend is a single-page application implemented in vanilla JavaScript without "
    "build tooling, transpilation or external dependencies. Layout is provided by CSS "
    "Flexbox and CSS Grid, with media queries adjusting padding and typography for "
    "narrower viewports. The user journey, described in Section 4.3, is implemented as "
    "three sequential views in the same HTML document: a landing screen, a three-step "
    "profile form and a plan view. View transitions are handled by toggling a small set "
    "of CSS classes rather than by a router, which keeps the JavaScript state minimal "
    "and the application robust to a forced refresh during testing."
)
add_figure(
    "IMG_8756.PNG",
    "Figure 5.1: CoachAI landing screen viewed on a phone. The value proposition (sport- and position-aware, equipment-adaptive, evidence-based) is presented above the fold; supported sports are surfaced as the first interactive element.",
    width_inches=2.5,
)
add_para(
    "The three-step profile form drives the recommender. Step one collects the user’s "
    "sport and playing position; the position list is filtered by the selected sport, so "
    "volleyball users see middle blocker, outside hitter, setter and libero, while "
    "basketball users see point guard, shooting guard, forward and centre (Figure 5.2). "
    "Step two collects the training goal (Figure 5.3) and the available equipment level. "
    "A short confirmation screen at step three echoes the selections back to the user "
    "before plan generation. The progress indicator at the top of the form (\"Step n of "
    "3\") is intended to reduce drop-off during onboarding, and the entire path can be "
    "completed in under a minute on a smartphone."
)
add_figure(
    "",
    "Figure 5.2: Sport-aware position pickers. Left: volleyball positions. Right: basketball positions. The position list is filtered by the sport selected immediately above, demonstrating the position-aware claim at the user-interface level.",
    width_inches=2.2,
    side_by_side=("IMG_8759.PNG", "IMG_8760.PNG"),
)
add_figure(
    "IMG_8761.PNG",
    "Figure 5.3: Goal selection (Step 2 of 3). The five training goals correspond directly to the GOAL_FOCUS lookup in the recommender (Section 5.3) and drive plan-level adjustments such as the additional plyometric block for the vertical-jump goal.",
    width_inches=2.5,
)
add_para(
    "On submission of the profile form, the frontend POSTs the four-element user profile "
    "to /api/generate-plan and renders the returned JSON as a list of session cards. "
    "Each session card shows the session name and focus, followed by a list of exercise "
    "blocks formatted as name, sets-by-reps, rest interval and a coaching note. The "
    "position-specific coaching note (Figure 5.4) is rendered prominently at the top of "
    "the plan view, immediately below the focus summary, so that the position-aware "
    "behaviour is visible to the user without further interaction."
)
add_figure(
    "IMG_8765.PNG",
    "Figure 5.4: Generated plan, top of the page. The header reflects the user’s sport and position (\"Weekly plan — Basketball (Point Guard)\"); the focus summary is drawn from GOAL_FOCUS; the highlighted position note is drawn from POSITION_NOTES (Section 5.3). The weekly progress indicator (D1–D5) supports FR7.",
    width_inches=2.5,
)
add_para(
    "Each session card is independently collapsible and carries a session-complete "
    "checkbox, which toggles a CSS class on the card and updates the small weekly progress "
    "indicator. Figure 5.5 shows the body of a generated plan; Figure 5.6 shows two "
    "sessions marked complete, with the corresponding checkboxes filled and the progress "
    "tracker updated. A print stylesheet hides the navigation chrome so that the plan can "
    "be printed or saved as a PDF for use in the gym (Figure 5.7), and an in-app feedback "
    "form at the foot of the plan view captures one-to-five-star ratings for clarity and "
    "sport-specificity, supporting FR8."
)
add_figure(
    "IMG_8766.PNG",
    "Figure 5.5: Generated plan, body of the page. The Lower-Body Power session (Day 1) closes with the Nordic hamstring curl and lateral bound; the Upper-Body Strength session (Day 2) follows with bench press, barbell row and accessory blocks.",
    width_inches=2.5,
)
add_figure(
    "IMG_8767.PNG",
    "Figure 5.6: Plan view with sessions marked complete. The Day 3 session card shows the green \"Done\" indicator after a session-complete tap, satisfying FR7. Day 4 (Power and Agility) follows below with sport-specific change-of-direction work.",
    width_inches=2.5,
)
add_figure(
    "IMG_8769.PNG",
    "Figure 5.7: Foot of the plan view. The print/save action is exposed alongside the in-app rating form, which captures one-to-five-star ratings for plan clarity and sport-specificity (FR8).",
    width_inches=2.5,
)

add_h2("5.5  Exercise Database and Extensibility")
add_para(
    "The exercise database is stored as a single JSON file, exercises.json, with one "
    "object per exercise and the schema described in Section 4.1.2. The forty-plus "
    "entries cover lower-body strength and power movements, plyometrics, upper-body push "
    "and pull patterns, core anti-rotation and bracing exercises, agility drills, and "
    "accessory exercises commonly prescribed for jump-sport athletes such as the Nordic "
    "hamstring curl and Copenhagen plank."
)
add_para(
    "Extending the system to a new sport or exercise is a data-only operation, satisfying "
    "the maintainability requirement (NFR6). Adding a new sport amounts to introducing "
    "the sport string to the supported list in generate_plan, adding a position lookup to "
    "POSITION_NOTES, and tagging existing exercises (or adding new ones) with the new "
    "sport in their sport_tags field. No changes are required to the filtering, picking "
    "or session-template logic. This separation between domain knowledge (the JSON data) "
    "and recommendation logic (the Python code) was a deliberate design decision and is "
    "the principal reason the implementation remained tractable within the project "
    "timeline."
)

add_h2("5.6  Cycle 2: Authentication and Per-User Plan Persistence")
add_para(
    "Cycle 2 introduces an optional account layer. Three new Python modules sit "
    "alongside the existing recommender: models.py defines the SQLAlchemy schema for the "
    "User and SavedPlan tables; auth.py exposes the signup, login, logout and current-"
    "user endpoints under the /api/auth prefix; and the existing app.py is extended with "
    "save-plan, list-plans, get-plan and per-plan progress-update endpoints. The "
    "additional dependencies, declared in requirements.txt, are Flask-SQLAlchemy and "
    "Flask-Login; both are well-supported and require no infrastructure beyond the SQLite "
    "file written into the backend directory at first run."
)
add_h3("5.6.1  Data Model")
add_para(
    "The data model has two tables. The User table records id, email (unique and "
    "indexed), the werkzeug-hashed password, an optional display name, and a creation "
    "timestamp. The SavedPlan table records id, the foreign key to its owning user, a "
    "user-supplied plan name, the four recommender inputs (sport, position, goal, "
    "equipment), the full plan JSON as a TEXT column, a JSON-encoded list of completed "
    "session indices, and a creation timestamp. The plan JSON is stored verbatim so that "
    "a saved plan survives later changes to the recommender or exercise database; the "
    "completion list is updated independently via the progress-update endpoint."
)
add_code(
    "class SavedPlan(db.Model):\n"
    "    __tablename__ = \"saved_plans\"\n"
    "    id = db.Column(db.Integer, primary_key=True)\n"
    "    user_id = db.Column(db.Integer, db.ForeignKey(\"users.id\"), nullable=False)\n"
    "    name = db.Column(db.String(120), nullable=False)\n"
    "    sport = db.Column(db.String(40), nullable=False)\n"
    "    position = db.Column(db.String(40), nullable=False)\n"
    "    goal = db.Column(db.String(40), nullable=False)\n"
    "    equipment = db.Column(db.String(40), nullable=False)\n"
    "    plan_json = db.Column(db.Text, nullable=False)\n"
    "    completed_sessions = db.Column(db.Text, default=\"[]\")\n"
    "    created_at = db.Column(db.DateTime, default=datetime.utcnow)"
)
add_h3("5.6.2  Authentication and Anonymous Use")
add_para(
    "The signup endpoint validates the email format, enforces a minimum password length "
    "of eight characters, rejects duplicate emails with HTTP 409, and on success creates "
    "a user, sets a session cookie via Flask-Login, and returns the new user payload. "
    "The login endpoint verifies the password using werkzeug’s constant-time check_"
    "password_hash and returns the same payload on success. Anonymous use is preserved "
    "throughout the rest of the API: the plan-generation endpoint never requires "
    "authentication, the home page does not require sign-in, and the save and list "
    "endpoints respond with a clean HTTP 401 if invoked without a session, which the "
    "frontend handles by routing the user through the sign-in flow."
)
add_h3("5.6.3  Save, List and Progress Endpoints")
add_para(
    "Plan saving is a single POST. The frontend submits the plan JSON together with a "
    "user-supplied name; the backend validates that the required keys are present, "
    "writes a SavedPlan row, and returns its summary. Listing returns the user’s plans "
    "ordered most-recent-first, including a derived completed_count field for the My "
    "Plans interface. The progress-update endpoint accepts a JSON list of session "
    "indices, validates that each is an integer in a sane range, deduplicates and sorts "
    "the list, and writes it as a JSON string to completed_sessions."
)
add_para(
    "A small backwards-compatibility helper at start-up, _ensure_schema, inspects the "
    "saved_plans table and adds the completed_sessions column if it is missing. This was "
    "introduced so that users who created a database before the progress feature was "
    "implemented can keep their existing data instead of having to drop and recreate the "
    "database. It is a deliberately narrow migration step rather than a full migration "
    "framework, which the project does not need."
)

add_h2("5.7  Cycle 2: LLM-Augmented Beginner-Friendly Notes")
add_para(
    "The clarity gap surfaced by the Phase 1 evaluation is addressed by an LLM-augmented "
    "subsystem with the architecture described in Section 4.2.2: the LLM is consulted "
    "only at development time, and the runtime system serves a static cache. The "
    "implementation has three parts: a build script, a runtime cache loader, and a small "
    "frontend toggle."
)
add_h3("5.7.1  Build Script")
add_para(
    "build_tiered_notes.py is a one-shot script that the developer runs locally. It "
    "reads the Gemini API key from the GEMINI_API_KEY environment variable, iterates "
    "over every exercise in exercises.json, and for each one issues a single Gemini call "
    "with a short system prompt that constrains the model strictly to rephrasing the "
    "existing coaching note. The system prompt enforces six explicit rules: no new "
    "exercises, no changes to volumes or numbers, no invented technique cues, a 40-word "
    "ceiling, plain everyday language with gym jargon spelled out, and output limited to "
    "the rephrased coaching note alone. The script writes the resulting strings to "
    "coaching_notes_tiered.json after every exercise so that a transient failure does "
    "not lose work, and supports resumption by skipping exercises that already appear in "
    "the cache."
)
add_h3("5.7.2  Runtime Cache Loader")
add_para(
    "At application start-up, recommender.py opens coaching_notes_tiered.json once and "
    "stores its contents in a module-level dictionary. If the file is missing the "
    "dictionary is empty and the system silently falls back to the original coaching "
    "note: the Cycle 1 behaviour is preserved exactly, which means that running CoachAI "
    "without ever generating the cache produces the same output as the Cycle 1 "
    "prototype. When the recommender is invoked with beginner_friendly=True, the _block "
    "helper that constructs each exercise block consults the cache and substitutes the "
    "beginner phrasing if one is present, leaving every other field of the block "
    "(category, sets, reps, rest) untouched."
)
add_h3("5.7.3  Frontend Toggle")
add_para(
    "The plan view in Cycle 2 carries a check-box toggle labelled Show beginner-friendly "
    "explanations. Ticking the toggle issues a fresh POST to /api/generate-plan with "
    "beginner_friendly set to true, and re-renders the returned plan in place. The "
    "toggle is off by default, which preserves the Cycle 1 baseline experience and "
    "supports the within-subjects A/B comparison reported in the Phase 2 evaluation "
    "(Chapter 6)."
)

add_h1("6.  Evaluation")
add_para(
    "This chapter reports the evaluation of CoachAI against the seven success criteria "
    "defined in Section 3.4. Reflecting the two-cycle development methodology described "
    "in Chapter 4, the user evaluation is conducted in two phases. Phase 1, reported in "
    "Section 6.3, evaluated the rule-based prototype produced in Cycle 1; its findings "
    "drove the scope of Cycle 2. Phase 2, reported in Section 6.4, is a smaller pilot of "
    "the extended system produced in Cycle 2. Functional testing (Section 6.2) addresses "
    "criteria 1, 2 and 3 and is reported once for the system as a whole. Section 6.5 "
    "records documentation and future-work evidence; Section 6.6 lists threats to "
    "validity; Section 6.7 summarises the outcome against each criterion."
)

add_h2("6.1  Evaluation Method")
add_para(
    "Three complementary methods are used, each chosen to match the kind of evidence the "
    "corresponding criteria require. Functional testing exhaustively exercises the "
    "recommender across the supported input space and inspects the resulting plan objects "
    "for structural validity. User evaluation is staged across two phases: Phase 1 "
    "collected perceived clarity and sport-specificity ratings from amateur athletes "
    "using the Cycle 1 rule-based prototype, while Phase 2 is a smaller pilot focused on "
    "the Cycle 2 additions (per-user persistence, the beginner-friendly toggle and "
    "progress tracking). Document inspection confirms that the requirements, design and "
    "implementation are recorded in this dissertation and that the future-work agenda is "
    "grounded in real evaluation findings rather than speculation. The combination "
    "provides triangulated evidence: system behaviour (functional testing), user "
    "perception (the two phases of user evaluation) and project deliverables "
    "(documentation)."
)

add_h2("6.2  Functional Evaluation (Criteria 1, 2 and 3)")
add_para(
    "A functional test matrix was generated by run_test_matrix.py, which iterates over "
    "every supported combination of sport, position, goal and equipment, calls the "
    "recommender for each combination, and validates the structure of the returned plan. "
    "The supported input space comprises two sports, four positions per sport, five "
    "training goals and three equipment levels, giving 120 distinct cases."
)
add_para(
    "Each plan is checked against three structural requirements: it must contain the "
    "expected top-level keys (sport, position, goal, equipment, focus_summary, "
    "position_note and sessions); it must include exactly five sessions; and every "
    "exercise block within each session must carry the fields name, category, sets, reps, "
    "rest and notes. Of the 120 input combinations, 120 produced a valid, well-structured "
    "plan, a 100% pass rate. Plans were generated successfully for both volleyball "
    "(60 of 60 cases) and basketball (60 of 60 cases). The full per-case results are "
    "reproduced in Appendix A."
)
add_para(
    "These results meet criteria 1 (every input combination produces a valid plan), 2 "
    "(at least two distinct sports supported) and 3 (plans are clearly structured) "
    "directly. The 100% pass rate is partly an artefact of the rule-based architecture: a "
    "deterministic system that consults a curated database is unlikely to produce "
    "structurally invalid output, because the failure modes that affect machine-learning "
    "or LLM-based recommenders (off-distribution prompts, hallucinated content, format "
    "drift) do not apply. The result is therefore strong evidence of structural reliability "
    "rather than of recommendation quality, which the user survey addresses."
)

add_h2("6.3  Phase 1: Rule-Based Prototype Evaluation")
add_h3("6.3.1  Design and Participants")
add_para(
    "Ten amateur athletes completed the survey instrument described in Appendix B: seven "
    "volleyball players and three basketball players. Two reported advanced gym experience "
    "(more than two years of consistent training), four reported intermediate experience "
    "(six months to two years), and four reported beginner experience (under six months). "
    "Recruitment targeted the Goldsmiths volleyball squad and a small number of basketball "
    "contacts. All responses were collected through Google Forms, anonymously and with "
    "explicit consent recorded as the first item in the survey. The instrument comprises "
    "three Likert items (mapped to criteria 4 and 5) and three open-ended free-text items, "
    "and the analysis below treats Likert responses as ordinal data, reporting the mean, "
    "the median and the proportion of responses scoring four or above. The system being "
    "evaluated in this phase is the Cycle 1 rule-based prototype: anonymous use, no "
    "persistence, no LLM-augmented coaching notes."
)

add_h3("6.3.2  Quantitative Results (Criteria 4 and 5)")
add_para(
    "Three Likert items, each on a scale from one (strongly disagree) to five (strongly "
    "agree), produced the headline numbers in Table 6.1. Question 4 measured perceived "
    "clarity of the plan and maps to criterion 4, which has a target of an average rating "
    "of at least four out of five. Questions 5 and 6 measured sport-specificity and map "
    "to criterion 5, which has a target of a majority of users agreeing or strongly "
    "agreeing (response of 4 or 5)."
)

# Likert results table
table_likert = doc.add_table(rows=1, cols=7)
table_likert.style = "Light Grid Accent 1"
hdr = table_likert.rows[0].cells
for i, t in enumerate(["Item", "Question", "Mean", "Median", "% ≥ 4", "Target", "Met?"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(10)
likert_rows = [
    ("Q4", "Plan is clearly structured and easy to follow", "3.30", "3.5", "50%", "mean ≥ 4 (criterion 4)", "No"),
    ("Q5", "Exercises feel relevant to my sport", "4.30", "4.0", "90%", "majority ≥ 4 (criterion 5)", "Yes"),
    ("Q6", "Plan is MORE sport-specific than my usual workouts", "3.30", "3.0", "40%", "majority ≥ 4 (criterion 5)", "No"),
]
for row in likert_rows:
    cells = table_likert.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
add_para("Table 6.1: Quantitative results for the three Likert items (N = 10).", italic=True)

add_para(
    "On Q4 (clarity), the mean rating of 3.30 falls below the criterion-4 target of 4.00. "
    "Five of ten participants rated the plan as four or above, and three rated it as two. "
    "This is an honest miss: criterion 4 was not met by the user-survey evidence collected. "
    "Functional structural validity is not in dispute (Section 6.2), but perceived clarity, "
    "which is what criterion 4 measures, is weaker than the project committed to."
)
add_para(
    "On Q5 (relevance to sport), the mean rating of 4.30 exceeds 4.00 and nine of ten "
    "participants rated relevance as four or higher. Q5 maps directly to one half of the "
    "criterion-5 framing and the result clearly meets it. On Q6, however, only four of ten "
    "participants rated CoachAI as more sport-specific than the workouts they usually do "
    "(mean 3.30, median 3.0). The criterion-5 target of a majority agreeing is therefore "
    "not met when the comparator is the user’s existing routine. Criterion 5 should be "
    "regarded as partially met: users find the exercises sport-relevant in absolute terms, "
    "but a majority do not yet judge CoachAI to outperform what they already do."
)

add_h3("6.3.3  Patterns by Experience and Sport")
add_para(
    "Disaggregating the Likert responses by experience level and by sport reveals two "
    "patterns that the headline numbers obscure. By experience, mean clarity ratings "
    "(Q4) rise monotonically from beginner (2.25) through intermediate (3.75) to advanced "
    "(4.50). The trend is consistent with the stated assumption in the project brief that "
    "users have basic gym experience and recognise common exercise names: when participants "
    "do not know the exercises by name, the bare name with sets and reps does not "
    "constitute a clear plan. Beginners are the group for whom criterion 4 fails most "
    "visibly. By contrast, perceived sport-specificity (Q6) does not follow the same "
    "monotonic trend: beginner mean 3.50, intermediate 3.25, advanced 3.00. Advanced users "
    "rate the plan as least different from their existing routines, consistent with their "
    "open-text responses indicating that they already follow self-directed programmes."
)
add_para(
    "By sport, volleyball participants rated clarity at 3.71 on average, while basketball "
    "participants rated it at 2.33. The basketball cohort is small (n = 3), so the "
    "comparison must be interpreted cautiously, but the direction is consistent with the "
    "implementation: the volleyball position notes, the volleyball-biased plyometric "
    "selections and the depth of volleyball-specific tagging in the exercise database all "
    "exceed the corresponding basketball coverage. Sport-relevance (Q5) was rated similarly "
    "by both cohorts (volleyball 4.29, basketball 4.33), suggesting that the sport-tag "
    "filtering is doing useful work even where the position-specific layer is thinner."
)

add_h3("6.3.4  Qualitative Themes from Free-Text Responses")
add_para(
    "Open-ended responses to Q7, Q8 and Q9 were read in full and grouped into themes. Five "
    "themes emerged with sufficient support to warrant reporting at this sample size."
)
add_para(
    "First, beginners and basketball users explicitly asked for exercise demonstrations or "
    "video instructions. Two participants requested either step-by-step instructions or a "
    "short animation showing how each exercise is performed. This theme aligns directly "
    "with the clarity gap observed in Q4: when users do not know an exercise by name, the "
    "current presentation does not bridge the gap. Second, three participants asked for a "
    "login system, progress tracking or the ability to save plans across sessions. This is "
    "a meaningful finding because each of these features was deliberately excluded from "
    "scope (Section 4.4), and the survey surfaces real user demand for them. Third, four "
    "participants asked for greater workout variety, particularly across different goals; "
    "one observed that the plans for general strength and vertical jump felt similar, which "
    "implies that the goal-driven adjustments are too subtle to be perceived. Fourth, two "
    "advanced participants reported that they would not use CoachAI in preference to their "
    "own routines, while several intermediate and beginner participants indicated they "
    "would use part or all of the plan. Fifth, one participant requested schedule "
    "flexibility, specifically the ability to add or remove training days to fit around "
    "academic commitments."
)

add_h2("6.4  Phase 2: Extended System Pilot")
add_h3("6.4.1  Pilot Design")
add_para(
    "Phase 2 is a small-N pilot focused on the additions delivered in Cycle 2: the "
    "authentication and per-user persistence layer, the saved-plans interface with "
    "progress tracking, and the LLM-augmented beginner-friendly coaching toggle. Three "
    "to five amateur athletes are recruited and asked to complete a guided task list "
    "covering the new features end to end: create an account, generate a plan, save it "
    "to their account, sign out, sign back in, reopen the plan from the My Plans list, "
    "mark one session complete, sign out and sign back in again, and confirm that the "
    "completion state has persisted. Participants are then asked to toggle the beginner-"
    "friendly explanations on and off and to compare the two presentations of the same "
    "exercise. Two short Likert items capture perceived clarity gain when the toggle is "
    "on (beginner-friendly notes felt clearer, 1–5) and ease-of-use of the saved-plans "
    "interface (1–5), supplemented by free-text feedback. The pilot is intentionally "
    "small: its purpose is to surface usability problems and qualitative impressions of "
    "the new features, not to produce statistically generalisable claims."
)
add_h3("6.4.2  Findings")
add_para(
    "[The Phase 2 pilot is scheduled for the week ending 16 May 2026. Findings will be "
    "reported in this section once the pilot is complete; the dissertation submission "
    "will include both quantitative and qualitative results for the new features.]",
    italic=True,
)

add_h2("6.5  Documentation and Future-Work Evidence (Criteria 6 and 7)")
add_para(
    "Criterion 6 requires that system design and implementation are fully documented. The "
    "present dissertation provides this evidence: Chapter 3 captures the requirements, "
    "Chapter 4 documents the design of both cycles, Chapter 5 walks through the "
    "implementation with code references for both cycles, and the appendices provide the "
    "functional test matrix, the survey instrument and supporting code listings. Criterion "
    "7 requires that future work is grounded in real evaluation findings. Cycle 2 itself "
    "is a direct response to Phase 1 findings — three of ten respondents asked for "
    "persistence and progress tracking; two asked for clearer instructions for novices; "
    "one observed that goal-driven plan differentiation felt too subtle. The remaining "
    "future-work agenda set out in Chapter 8 is similarly grounded in evaluation findings "
    "rather than on speculative extensions."
)

add_h2("6.6  Threats to Validity")
add_para(
    "Several threats to validity apply across both phases. First, sample sizes are small "
    "and skewed: Phase 1 recruited N = 10, dominated by volleyball participants and "
    "beginner-to-intermediate experience levels; Phase 2 is by design a small pilot. The "
    "Phase 1 results should be treated as small-N evidence and replication with a larger, "
    "more balanced cohort is required before the headline ratings can be generalised. "
    "Second, recruitment in both phases was through the developer’s own social network, "
    "which introduces a possible courtesy bias: participants may be more positive than a "
    "neutral cohort, particularly on items asking how relevant the plan feels. Third, the "
    "Phase 1 evaluation captured first-impression responses rather than the experience of "
    "following a plan in the gym over weeks, and Phase 2 is similarly cross-sectional. The "
    "behaviours that would matter most for a longitudinal analysis — does saved-plan "
    "persistence change adherence, does the LLM-augmented beginner toggle improve "
    "clarity over multiple sessions — are out of reach of this study and identified as "
    "future work in Chapter 8."
)

add_h2("6.7  Summary Against the Success Criteria")
add_para(
    "Table 6.2 summarises the outcome against each criterion based on Phase 1 evidence. "
    "Three of the seven criteria are met fully on functional grounds (1, 2, 3); two are "
    "met on the basis of this dissertation and the explicit grounding of Cycle 2 in "
    "Phase 1 findings (6 and 7); one is partially met by Phase 1 (criterion 5: relevance "
    "is met, more-specific-than-usual is not); and one is not met by Phase 1 alone "
    "(criterion 4: perceived clarity below target). The Cycle 2 additions, particularly "
    "the beginner-friendly explanations toggle, were designed specifically to address the "
    "Phase 1 miss on criterion 4; the Phase 2 pilot will report on whether this addition "
    "narrows the clarity gap. The misses are reported faithfully and inform the "
    "discussion in Chapter 7."
)

# Final summary table
table_summary = doc.add_table(rows=1, cols=4)
table_summary.style = "Light Grid Accent 1"
hdr = table_summary.rows[0].cells
for i, t in enumerate(["#", "Criterion", "Outcome", "Evidence"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(10)

summary_rows = [
    ("1", "Valid plan from any input combination", "Met", "120/120 functional test cases passed (Section 6.2)"),
    ("2", "At least two sports supported", "Met", "60/60 volleyball and 60/60 basketball cases passed"),
    ("3", "Plans clearly structured", "Met", "100% structural validity across the test matrix"),
    ("4", "Plans easy to understand (mean ≥ 4/5)", "Not met", "Q4 mean 3.30 (target 4.00); 50% rated ≥ 4"),
    ("5", "Plans sport-specific vs generic (majority ≥ 4)", "Partially met", "Q5 relevance 90% ≥ 4 (met); Q6 more-specific 40% ≥ 4 (not met)"),
    ("6", "System fully documented", "Met", "Chapters 3, 4, 5 plus Appendices A–C"),
    ("7", "Future work grounded in evaluation findings", "Met", "Chapter 8 grounded in this chapter’s misses and themes"),
]
for row in summary_rows:
    cells = table_summary.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
add_para("Table 6.2: Outcome against each of the seven project success criteria.", italic=True)

add_h1("7.  Discussion and Reflection")
add_para(
    "This chapter discusses the evaluation findings of Chapter 6 in four parts. Section "
    "7.1 examines what CoachAI does well; Section 7.2 examines where it falls short and "
    "why; Section 7.3 reflects on the iterative two-cycle process, examining how the "
    "Phase 1 findings shaped Cycle 2 and what the development sequence cost and gained; "
    "and Section 7.4 records methodological limitations of the study itself."
)

add_h2("7.1  Where CoachAI Performs Well")
add_para(
    "Three results from the evaluation are worth dwelling on. First, structural reliability "
    "is high and unambiguous: the recommender produces a valid, well-formed plan for every "
    "supported input combination, and the 100% pass rate held across all three equipment "
    "levels and both sports. The recommender’s deterministic, rule-based architecture "
    "makes this property easy to obtain and easy to argue for, and it gives users a stable "
    "artefact they can trust on every visit. Second, sport-relevance ratings on Q5 were "
    "strong (mean 4.30, 90% of participants rating four or higher), and were consistent "
    "across both sports. This suggests that the relatively simple sport-tagging mechanism "
    "in the exercise database is doing useful work: when a volleyball middle blocker "
    "receives broad jumps, depth jumps and the approach jump in Day 1, those exercises "
    "register as sport-relevant to the participants reading the plan. Third, the experience-"
    "level pattern (clarity rising from beginner 2.25 to advanced 4.50) is consistent with "
    "the project’s stated assumption that users have basic gym experience and recognise "
    "common exercise names; in other words, the system is performing as designed for the "
    "audience it explicitly targeted, while also revealing where that target needs to be "
    "broadened."
)

add_h2("7.2  Where CoachAI Falls Short")
add_para(
    "Two findings demand a more critical reading. The first is that perceived clarity is "
    "below the criterion-4 target, with a particularly steep gap for beginners (mean 2.25). "
    "The free-text responses make the cause concrete: beginners and basketball participants "
    "asked for exercise demonstrations or video instructions. The current implementation "
    "displays an exercise as a name, sets and repetitions, rest interval and a one-sentence "
    "coaching note. For a participant who has never performed a Bulgarian split squat or a "
    "depth jump, this is insufficient even when the prescription itself is appropriate. "
    "The clarity miss is therefore not a failure of recommendation logic but a failure of "
    "presentation: the right exercise expressed in the wrong medium for novice users."
)
add_para(
    "The second is that, while users find the exercises sport-relevant in absolute terms "
    "(Q5), only a minority feel that CoachAI is more sport-specific than the workouts they "
    "already do (Q6 mean 3.30, 40% rating four or higher). This is most visible for "
    "advanced participants, two of whom reported having their own routines and would not "
    "switch to CoachAI in preference. This pattern is internally consistent and, in some "
    "ways, expected: advanced users have invested time in personal programmes and are "
    "unlikely to be displaced by a first-impression interaction with a prototype. The more "
    "informative observation is that even intermediate participants rated Q6 modestly "
    "(mean 3.25), which suggests that the differentiation between CoachAI’s plans and "
    "available generic gym programmes, as perceived by the user, is real but not "
    "overwhelming. Two specific hooks for improvement emerge. First, the goal-driven "
    "adjustments, currently limited to an extra plyometric block for the vertical-jump "
    "goal, are too subtle to be perceived: one participant explicitly noted that the "
    "general-strength and vertical-jump plans felt similar. Second, the basketball cohort "
    "rated clarity considerably below volleyball (2.33 vs 3.71), which is consistent with "
    "the implementation: the position notes, plyometric variants and exercise tagging are "
    "deeper for volleyball than for basketball, and broadening the basketball coverage is "
    "an obvious next step."
)

add_h2("7.3  Reflecting on the Iterative Process")
add_para(
    "The two-cycle structure described in Chapter 4 is more than a presentation device: "
    "it shaped what got built and in what order. Cycle 1 deliberately delivered a lean, "
    "stateless rule-based prototype, prioritising structural reliability and a tractable "
    "evaluation. Phase 1 of the user study then surfaced two patterns that Cycle 2 was "
    "scoped to address. Three of ten participants asked for some form of plan saving or "
    "progress tracking; two asked for richer instructions for novice users; and the "
    "experience-level breakdown of clarity ratings (beginner 2.25, advanced 4.50) "
    "indicated that the clarity gap was concentrated in beginners. Cycle 2 responded with "
    "an authentication-and-persistence layer (saved plans with per-session progress "
    "tracking) and an LLM-augmented beginner-friendly notes toggle — both features tied "
    "directly to specific Phase 1 findings rather than to speculation about what users "
    "might want."
)
add_para(
    "This iterative shape paid two dividends. First, it allowed structural-reliability "
    "claims to be made cleanly: the 100% functional pass rate reported in Section 6.2 was "
    "established against the deterministic Cycle 1 system, and Cycle 2’s additions were "
    "designed to leave that property intact (the LLM is consulted only at development "
    "time; the runtime uses a static cache; authentication does not change the "
    "recommender output for a given input). Second, it grounded the dissertation’s "
    "future-work framing differently from how it would have read had everything been "
    "delivered in one cycle. Chapter 8’s remaining future-work items are not mid-project "
    "feature requests that ran out of time, but extensions that the two-phase evaluation "
    "actively did not test."
)
add_para(
    "An honest cost of the iterative shape is that the Phase 2 evaluation is necessarily "
    "smaller than Phase 1 and is reported as a pilot rather than as a comparable user "
    "study. With more time, a balanced two-phase study with similar N in each phase would "
    "permit a stronger claim about whether Cycle 2 narrowed the clarity gap. As reported, "
    "the Phase 2 pilot’s qualitative findings should be read as directional evidence about "
    "the new features rather than as definitive validation of the rephrasing layer."
)

add_h2("7.4  Methodological Reflection")
add_para(
    "Three methodological points deserve explicit acknowledgement. First, the user-survey "
    "sample of ten is appropriate for a final-year prototype but cannot support strong "
    "generalisation. The findings are directionally informative — the magnitudes and the "
    "experience-level pattern are clear — but they should be replicated with a larger and "
    "more balanced cohort, particularly more basketball participants and more advanced "
    "athletes. Second, the evaluation captures first-impression responses rather than "
    "longitudinal behaviour: participants generated a plan, read it and rated it, but did "
    "not follow it in the gym over a multi-week period. The behaviours that matter most "
    "for criteria 4 and 5 — does the plan remain clear after seven sessions, does it "
    "produce perceived progress over weeks — are out of reach of this study and identified "
    "as future work in Chapter 8. Third, the recruitment route, which leaned on the "
    "developer’s social network, may introduce a courtesy bias that inflates Q5 ratings; "
    "the comparatively flat Q6 distribution suggests this bias has not displaced the "
    "more critical question, but it is worth recording."
)
add_para(
    "Reflecting on the development process itself, the most consequential decision was the "
    "early move to a stateless rule-based architecture grounded in a curated exercise "
    "database. That choice constrained scope in a way that allowed the project to deliver "
    "a demonstrable, testable system within the available time, and it is the same choice "
    "that made the structural-validity claim of Chapter 6 tractable. The cost, made visible "
    "by the survey, is that the system as built is best for the user it explicitly "
    "targeted — an intermediate amateur athlete with basic gym vocabulary — and is weaker "
    "for the user populations on either side of that profile. Acknowledging this honestly "
    "is more useful than claiming a generality the evaluation does not support."
)

add_h1("8.  Conclusion and Future Work")
add_para(
    "This chapter closes the dissertation by restating the contribution (Section 8.1), "
    "setting out a future-work agenda grounded in the evaluation findings (Section 8.2), "
    "and offering brief closing remarks (Section 8.3)."
)

add_h2("8.1  Summary of Contribution")
add_para(
    "CoachAI is a mobile-first web application that generates structured weekly strength "
    "and conditioning plans for amateur athletes from four user inputs (sport, playing "
    "position, training goal and available equipment), grounded in published S&C "
    "literature and supporting volleyball and basketball as initial case studies. The "
    "system was developed in two cycles. Cycle 1 delivered a transparent rule-based "
    "recommender backed by a curated database of more than forty exercises, evaluated "
    "with ten amateur athletes whose feedback shaped the next iteration. Cycle 2 added "
    "an optional account layer with per-user plan persistence and progress tracking, "
    "and an LLM-augmented beginner-friendly coaching toggle implemented as a runtime "
    "static cache authored at development time. Across the seven success criteria "
    "committed to at project inception, five are met fully (criteria 1, 2, 3, 6 and 7), "
    "one is partially met (criterion 5), and one was not met by the Phase 1 evaluation "
    "(criterion 4: perceived clarity). The Cycle 2 additions were scoped specifically to "
    "address the Phase 1 misses; the small Phase 2 pilot reports on whether they do."
)

add_h2("8.2  Future Work")
add_para(
    "The evaluation findings translate into a focused, evidence-based future-work agenda. "
    "Each item below is grounded in a specific finding from Chapter 6, and the items that "
    "Cycle 2 has already delivered (per-user persistence and progress tracking; an LLM-"
    "augmented beginner-friendly coaching toggle) are recorded in Section 8.1 rather "
    "than restated as future work."
)
add_para(
    "Exercise demonstrations beyond text rephrasing. The Cycle 2 LLM toggle addresses the "
    "clarity gap at the level of plain-language instructions; for some beginners, a "
    "static illustrated card, a looping animation or a short embedded video would be "
    "more effective still. Two Phase 1 participants requested visual demonstrations; "
    "this would build directly on the existing toggle interface and could be added "
    "without changing the recommender."
)
add_para(
    "Richer goal differentiation. One Phase 1 participant noted that the general-strength "
    "and vertical-jump plans felt similar, indicating that the goal-driven adjustments "
    "are too subtle to be perceived. A more substantial set of goal-specific session "
    "templates (for example, separate templates for strength, hypertrophy, vertical jump "
    "and agility goals, rather than templates plus a single additional plyometric block) "
    "would produce more visibly differentiated outputs and address the Q6 miss directly."
)
add_para(
    "Broader basketball coverage. The basketball cohort rated clarity considerably below "
    "volleyball (2.33 vs 3.71). Although the basketball sample is small (n = 3), the "
    "direction is consistent with the implementation, where volleyball-specific tagging, "
    "position notes and exercise variants are deeper. Adding more basketball-specific "
    "exercises and richer position notes is a data-only change (NFR6 maintainability) "
    "that should narrow the gap."
)
add_para(
    "Schedule flexibility. One Phase 1 participant asked for the ability to add or "
    "remove training days to fit around academic commitments. A future iteration could "
    "allow users to customise the number of weekly sessions, with the recommender "
    "re-balancing category coverage to preserve programme integrity."
)
add_para(
    "Longitudinal evaluation. Both phases of the user study captured short-horizon "
    "responses rather than the experience of following a plan in the gym over weeks. A "
    "multi-week diary study with a smaller cohort would better answer the durable-"
    "clarity and perceived-progress questions that criteria 4 and 5 implicitly assume, "
    "and would specifically test whether the Cycle 2 features (saved plans, progress "
    "tracking, beginner-friendly notes) change adherence over time."
)
add_para(
    "Wider sport coverage. The system was deliberately designed so that adding a sport is "
    "a data-only change. Demonstrating this empirically with a third sport (for example "
    "football or netball) would test the maintainability requirement (NFR6) and provide "
    "stronger evidence of the multi-sport claim than two sports alone."
)
add_para(
    "Native mobile applications. The current system is a mobile-first responsive web "
    "application, which was the right scope for the project but produces a different user "
    "experience to a native iOS or Android application (no home-screen icon, no push "
    "notifications, no deep platform integration). Wrapping the existing API and "
    "frontend in a native shell, or porting to a cross-platform framework, would be a "
    "natural extension and is the only feature originally placed out of scope by the "
    "project brief that has not been at least partially addressed."
)

add_h2("8.3  Closing Remarks")
add_para(
    "The honest summary is that CoachAI demonstrates that a transparent, rule-based "
    "recommender can produce structurally reliable, sport-relevant strength and "
    "conditioning plans for amateur athletes, and that an iterative two-cycle design — "
    "prototype, evaluate with users, extend on the basis of what the evaluation surfaces "
    "— is a feasible methodology for delivering and reporting a final-year project of "
    "this scope. The Cycle 1 design choices (statelessness, determinism, explicit rules) "
    "held up under evaluation; the Cycle 2 additions (per-user persistence, progress "
    "tracking, an LLM-augmented beginner-friendly toggle) were each scoped to address "
    "specific Phase 1 findings, and were implemented in ways that preserved the Cycle 1 "
    "properties rather than displacing them. The system retains clear weaknesses for "
    "cohorts with established personal routines and for very novice users beyond the "
    "reach of text-only rephrasing; both are recorded in the future-work agenda. Both "
    "the strengths and the weaknesses are reported faithfully, and what comes next is "
    "built on what the data actually says."
)


# ---------------------------------------------------------------------------
# REFERENCES
# ---------------------------------------------------------------------------

add_h1("References")
refs = [
    "Abramson, A. (2025) ‘Fitbit’s personal health coach in public preview is here’, Fitbit / Google Blog, 27 October. Available at: https://blog.google/products/fitbit/personal-health-coach-public-preview/ (Accessed: 3 November 2025).",
    "AI Endurance (2024a) AI Endurance: AI running, cycling, and triathlon coach. Available at: https://aiendurance.com/en (Accessed: 28 September 2025).",
    "AI Endurance (2024b) How does AI Endurance work? Available at: https://aiendurance.com/en/product (Accessed: 7 October 2025).",
    "Cao, J., Xu, P., Zhang, L. and Liu, Q. (2022) ‘Hybrid recommendation models in personalised health and fitness systems’, ACM Computing Surveys, 55(8), pp. 1–29.",
    "Farah, M. and Clarke, A. (2025) ‘4x Olympic Gold Medalist Sir Mo Farah launches URUNN running app’, Fitt Insider – Press Release. Available at: https://insider.fitt.co/press-release/4x-olympic-gold-medalist-sir-mo-farah-launches-urunn-running-app/ (Accessed: 10 October 2025).",
    "Felfernig, A., Friedrich, G. and Jannach, D. (2014) Recommender Systems: An Introduction. Cambridge: Cambridge University Press.",
    "FitnessAI (2025a) FitnessAI — Get stronger with A.I. Available at: https://www.fitnessai.com/ (Accessed: 1 November 2025).",
    "FitnessAI (2025b) FitnessAI: Gym Workout Planner (App Store / Google Play descriptions). Available at: https://apps.apple.com/us/app/fitness-ai-gym-workout-planner/id1446224156 (Accessed: 14 September 2025).",
    "Google/Fitbit (2025) Fitbit unveils Gemini-powered personal health coach. Available at: https://blog.google/products/fitbit/ (Accessed: 13 November 2025).",
    "Grădinaru, L. (2021) ‘Plyometric training effectiveness on vertical jump in junior female volleyball players: a systematic review’, Discobolul – Physical Education, Sport and Kinetotherapy Journal, 17(3), pp. 120–132.",
    "Ma, S., Li, H., Giannakis, M., Tang, Z. and Huang, Y. (2025) ‘Effects of physical training programs on healthy athletes’ vertical jump performance: a systematic review and meta-analysis’, Journal of Sports Science and Medicine, 24(3), pp. 236–247.",
    "Markovic, G. (2007) ‘Does plyometric training improve vertical jump height? A meta-analytical review’, British Journal of Sports Medicine, 41(6), pp. 349–355.",
    "Newton, R.U. and Dugan, E. (2002) ‘Application of strength diagnosis’, Strength and Conditioning Journal, 24(5), pp. 50–59.",
    "Ramirez-Campillo, R., Andrade, D.C., Nikolaidis, P.T., Moran, J., Clemente, F.M., Chaabene, H. and Comfort, P. (2020) ‘Effects of plyometric jump training on vertical jump height of volleyball players: a systematic review with meta-analysis of randomized-controlled trials’, Journal of Sports Science and Medicine, 19(3), pp. 489–499.",
    "Ricci, F., Rokach, L. and Shapira, B. (2015) Recommender Systems Handbook. 2nd edn. New York: Springer.",
    "Sheppard, J.M. and Gabbett, T. (2009) ‘An analysis of the relative importance of strength and power qualities in volleyball’, Journal of Strength and Conditioning Research, 23(7), pp. 2050–2056.",
    "WithU (2025) ‘WithU partners with Sir Mo Farah to launch global running app’, Ecommerce News UK, 2 May. Available at: https://ecommercenews.uk/story/withu-partners-with-sir-mo-farah-to-launch-global-running-app (Accessed: 12 November 2025).",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# APPENDICES (placeholders)
# ---------------------------------------------------------------------------

add_h1("Appendices")

add_h2("Appendix A: Functional Test Matrix")
add_para(
    "The functional test matrix was generated by run_test_matrix.py, which iterates over "
    "every supported combination of (sport, position, goal, equipment) and validates each "
    "resulting plan for structural correctness. The full output is provided as supporting "
    "files alongside this dissertation: appendix_a_test_matrix.md (human-readable) and "
    "appendix_a_test_matrix.csv (machine-readable)."
)
add_para("Headline result, summarised in Section 6.2:", bold=True)
add_bullet("Total cases tested: 120 (2 sports × 4 positions × 5 goals × 3 equipment levels)")
add_bullet("Cases passing structural validation: 120 of 120 (100.0%)")
add_bullet("Volleyball cases passing: 60 of 60")
add_bullet("Basketball cases passing: 60 of 60")
add_bullet("Failures: 0")

add_h2("Appendix B: User Survey Instrument")
add_para(
    "The survey instrument used for the user evaluation reported in Section 6.3 is provided "
    "as a separate document, user_survey.docx, alongside this dissertation. It includes the "
    "participant information sheet, the consent statement, the demographic questions, the "
    "three Likert items mapped to criteria 4 and 5, and the three open-ended free-text "
    "items. The raw anonymised responses are provided as survey_responses.csv."
)

add_h2("Appendix C: Selected Code Listings and Extended Screenshots")
add_para(
    "The full source code is provided alongside this dissertation in the coachai/ folder. "
    "The most important files for the recommender are listed below; selected extracts are "
    "reproduced in Chapter 5."
)
add_bullet("backend/app.py — Flask application and REST endpoint (≈53 lines).")
add_bullet("backend/recommender.py — Rule-based plan generator with the equipment-hierarchy filter, deterministic picker and five session templates (≈239 lines).")
add_bullet("backend/exercises.json — Curated exercise database with 40+ entries (sport tags, equipment level, category, coaching notes).")
add_bullet("frontend/index.html — Single-page application HTML structure.")
add_bullet("frontend/app.js — Profile form, fetch call to /api/generate-plan, plan rendering.")
add_bullet("frontend/styles.css — Mobile-first responsive stylesheet.")
add_para(
    "The following supplementary figure shows a session-complete state midway through the "
    "weekly plan, with Day 4 marked complete and Day 5 (Core and Recovery) starting "
    "directly below."
)
add_figure(
    "IMG_8768.PNG",
    "Figure C.1: Plan view, Day 4 marked complete. The green \"Done\" indicator and the populated weekly progress tracker (cf. Figure 5.4) provide visible state-change feedback to the user.",
    width_inches=2.5,
)


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

OUT_PATH = os.path.join(os.path.dirname(__file__), "dissertation.docx")
doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
